"""
Genera el documento de entrega del proyecto en formato Word.
Al respecto, el documento reúne en un solo archivo todo lo que pide la consigna, con las secciones ordenadas y con la evidencia de ejecución tomada directamente del reporte que dejó la última corrida.
Dicho detalle importa, porque los números que aparecen en el documento no están escritos a mano sino leídos del resultado real, de modo que no pueden quedar desactualizados respecto del código.

Uso::

    python scripts/generar_documento_entrega.py
    python scripts/generar_documento_entrega.py --destino /otra/ruta.docx
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parent.parent
if str(RAIZ) not in sys.path:
    sys.path.insert(0, str(RAIZ))

# La paleta del documento se declara una sola vez para que todas las secciones tengan el mismo aspecto.
# De ese modo, cambiar el color de los encabezados o el de las notas es tocar un solo renglón y no recorrer el archivo entero.
AZUL_PRINCIPAL = RGBColor(0x1F, 0x3A, 0x5F)
AZUL_SECUNDARIO = RGBColor(0x2E, 0x5E, 0x8C)
GRIS_TEXTO = RGBColor(0x33, 0x33, 0x33)
GRIS_SUAVE = RGBColor(0x66, 0x66, 0x66)
VERDE = RGBColor(0x1E, 0x7B, 0x3C)

DESTINO_POR_DEFECTO = RAIZ / "Entrega_Final_Pipeline_Ventas_Minoristas.docx"


# =============================================================================
# Utilidades de formato
# =============================================================================


def _sombrear(elemento, color_hexadecimal: str) -> None:
    """
    Aplica un color de fondo a una celda de tabla o a un párrafo.
    Cabe señalar que python-docx no expone el sombreado, motivo por el cual hay que insertar el elemento XML a mano dentro de las propiedades del objeto.
    Recibe el elemento a sombrear y el color expresado en seis dígitos hexadecimales sin almohadilla delante.
    """
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:color"), "auto")
    sombreado.set(qn("w:fill"), color_hexadecimal)

    propiedades = elemento._element.get_or_add_tcPr() if hasattr(
        elemento._element, "get_or_add_tcPr"
    ) else elemento._p.get_or_add_pPr()
    propiedades.append(sombreado)


def _borde_inferior(parrafo, color_hexadecimal: str, grosor: int = 12) -> None:
    """
    Dibuja una línea debajo de un párrafo.
    Recibe el párrafo al que agregarle la línea, el color de esa línea y su grosor.
    Conviene precisar que el grosor se expresa en octavos de punto, que es la unidad con la que trabaja el formato de Word.
    """
    propiedades = parrafo._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    inferior = OxmlElement("w:bottom")
    inferior.set(qn("w:val"), "single")
    inferior.set(qn("w:sz"), str(grosor))
    inferior.set(qn("w:space"), "4")
    inferior.set(qn("w:color"), color_hexadecimal)
    bordes.append(inferior)
    propiedades.append(bordes)


def configurar_estilos(documento: Document) -> None:
    """
    Define la tipografía y el espaciado de todo el documento.
    Recibe el documento sobre el que aplicar los estilos y lo modifica en el lugar, tanto en el texto corrido como en los cuatro niveles de encabezado.
    """
    normal = documento.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = GRIS_TEXTO
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    # La fuente asiática se fija aparte porque, de lo contrario, Word toma otra por defecto y los acentos pueden aparecer con un tipo distinto al del resto del renglón.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for nivel, tamanio, color in (
        (1, 20, AZUL_PRINCIPAL),
        (2, 15, AZUL_SECUNDARIO),
        (3, 12.5, AZUL_SECUNDARIO),
        (4, 11.5, GRIS_TEXTO),
    ):
        estilo = documento.styles[f"Heading {nivel}"]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(tamanio)
        estilo.font.bold = True
        estilo.font.color.rgb = color
        estilo.paragraph_format.space_before = Pt(16 if nivel <= 2 else 12)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.keep_with_next = True


def titulo(documento: Document, texto: str, nivel: int = 1, salto: bool = False) -> None:
    """
    Escribe un encabezado con el nivel indicado.
    Recibe el documento donde escribir, el texto del encabezado y su nivel, que va de uno a cuatro.
    En caso de que se pida un salto, el encabezado empieza en una página nueva.
    Los encabezados de primer nivel llevan además una línea inferior, que es lo que separa visualmente las secciones principales.
    """
    if salto:
        documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    encabezado = documento.add_heading(texto, level=nivel)
    if nivel == 1:
        _borde_inferior(encabezado, "1F3A5F", 16)


def parrafo(documento: Document, texto: str, cursiva: bool = False) -> None:
    """
    Escribe un párrafo de texto corrido y justificado.
    Recibe el documento donde escribir y el contenido del párrafo.
    En caso de que se pida cursiva, el texto se pinta además en gris suave, que es el tratamiento reservado a las aclaraciones.
    """
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corrida = p.add_run(texto)
    corrida.italic = cursiva
    if cursiva:
        corrida.font.color.rgb = GRIS_SUAVE


def vinetas(documento: Document, elementos: list[str], numeradas: bool = False) -> None:
    """
    Escribe una lista con viñetas o numerada.
    Recibe el documento donde escribir y el texto de cada elemento de la lista.
    En caso de que se pida una lista numerada, se aplica el estilo correspondiente en lugar del de viñetas.
    Conviene precisar que lo que va entre un par de asteriscos se resalta en negrita al pintarlo, de modo que los datos no necesitan traer marcado dentro.
    """
    estilo = "List Number" if numeradas else "List Bullet"
    for elemento in elementos:
        p = documento.add_paragraph(style=estilo)
        p.paragraph_format.space_after = Pt(3)
        # Al partir por el par de asteriscos, los fragmentos de posición impar son justamente los que estaban encerrados entre ellos, razón por la que reciben negrita.
        for indice, fragmento in enumerate(elemento.split("**")):
            corrida = p.add_run(fragmento)
            corrida.bold = indice % 2 == 1


def codigo(documento: Document, contenido: str, lenguaje: str = "") -> None:
    """
    Inserta un bloque de código con fondo gris y tipografía monoespaciada.
    Al respecto, se usa una tabla de una sola celda en lugar de un párrafo sombreado, puesto que de ese modo el bloque conserva su fondo aunque el texto ocupe varias líneas.
    Recibe el documento donde escribir y el código a mostrar.
    Adicionalmente admite una etiqueta de lenguaje, que se imprime encima del bloque cuando se indica.
    """
    if lenguaje:
        etiqueta = documento.add_paragraph()
        etiqueta.paragraph_format.space_after = Pt(2)
        corrida = etiqueta.add_run(lenguaje)
        corrida.font.size = Pt(8)
        corrida.font.bold = True
        corrida.font.color.rgb = GRIS_SUAVE

    tabla = documento.add_table(rows=1, cols=1)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    celda = tabla.cell(0, 0)
    _sombrear(celda, "F4F5F7")

    celda.text = ""
    for indice, linea in enumerate(contenido.rstrip("\n").split("\n")):
        p = celda.paragraphs[0] if indice == 0 else celda.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        corrida = p.add_run(linea if linea else " ")
        corrida.font.name = "Consolas"
        corrida.font.size = Pt(8.5)
        corrida.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        corrida.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")

    documento.add_paragraph().paragraph_format.space_after = Pt(4)


def tabla(
    documento: Document,
    encabezados: list[str],
    filas: list[list[str]],
    anchos: list[float] | None = None,
) -> None:
    """
    Inserta una tabla con la fila de encabezado resaltada.
    Recibe el documento donde escribir, los textos del encabezado y el contenido de cada una de las filas.
    Adicionalmente admite el ancho de cada columna expresado en centímetros.
    Las filas de posición impar se sombrean en un gris muy claro, con el fin de que la vista siga el renglón en las tablas largas.
    """
    t = documento.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    fila_encabezado = t.rows[0]
    for indice, texto in enumerate(encabezados):
        celda = fila_encabezado.cells[indice]
        _sombrear(celda, "1F3A5F")
        celda.text = ""
        p = celda.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        corrida = p.add_run(texto)
        corrida.bold = True
        corrida.font.size = Pt(9.5)
        corrida.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for numero, contenido in enumerate(filas):
        fila = t.add_row()
        if numero % 2 == 1:
            for celda in fila.cells:
                _sombrear(celda, "F7F8FA")
        for indice, texto in enumerate(contenido):
            celda = fila.cells[indice]
            celda.text = ""
            p = celda.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            corrida = p.add_run(str(texto))
            corrida.font.size = Pt(9.5)

    if anchos:
        for fila in t.rows:
            for indice, ancho in enumerate(anchos):
                if indice < len(fila.cells):
                    fila.cells[indice].width = Cm(ancho)

    documento.add_paragraph().paragraph_format.space_after = Pt(4)


def nota(documento: Document, texto: str, tipo: str = "info") -> None:
    """
    Inserta un recuadro destacado.
    Recibe el documento donde escribir y el contenido de la nota.
    El tipo determina el color del recuadro y admite los valores "info", "exito" y "aviso".
    Al igual que en las listas, lo que va entre un par de asteriscos se resalta en negrita.
    """
    colores = {"info": "E8F0FA", "exito": "E8F5EC", "aviso": "FDF3E3"}
    bordes = {"info": AZUL_SECUNDARIO, "exito": VERDE, "aviso": RGBColor(0xB0, 0x6D, 0x00)}

    t = documento.add_table(rows=1, cols=1)
    celda = t.cell(0, 0)
    _sombrear(celda, colores.get(tipo, "E8F0FA"))
    celda.text = ""
    p = celda.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    for indice, fragmento in enumerate(texto.split("**")):
        corrida = p.add_run(fragmento)
        corrida.font.size = Pt(10)
        corrida.bold = indice % 2 == 1
        corrida.font.color.rgb = bordes.get(tipo, AZUL_SECUNDARIO)

    documento.add_paragraph().paragraph_format.space_after = Pt(4)


# =============================================================================
# Secciones del documento
# =============================================================================


def portada(documento: Document, datos: dict) -> None:
    """
    Arma la primera página del documento.
    Recibe el documento donde escribir y el diccionario con los metadatos de la entrega, del que toma el autor, el repositorio, la fecha y la versión.
    """
    for _ in range(4):
        documento.add_paragraph()

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = p.add_run("PROYECTO INTEGRADOR DE INGENIERÍA DE DATOS")
    corrida.font.size = Pt(11)
    corrida.font.bold = True
    corrida.font.color.rgb = GRIS_SUAVE

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = p.add_run("Pipeline de datos extremo a extremo\nsobre ventas minoristas")
    corrida.font.size = Pt(28)
    corrida.font.bold = True
    corrida.font.color.rgb = AZUL_PRINCIPAL

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _borde_inferior(p, "2E5E8C", 20)

    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corrida = p.add_run(
        "Ingesta, validación de calidad, transformación, modelado,\n"
        "procesamiento distribuido y observabilidad"
    )
    corrida.font.size = Pt(13)
    corrida.font.color.rgb = GRIS_SUAVE
    corrida.italic = True

    for _ in range(5):
        documento.add_paragraph()

    t = documento.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    entradas = [
        ("Autor", datos["autor"]),
        ("Repositorio", datos["repositorio"]),
        ("Conjunto de datos", "Online Retail II, histórico real de 1,067,371 transacciones"),
        ("Período cubierto", "1 de diciembre de 2009 al 9 de diciembre de 2011"),
        ("Fecha de entrega", datos["fecha"]),
        ("Versión", datos["version"]),
    ]

    for etiqueta, valor in entradas:
        fila = t.add_row()
        celda_etiqueta = fila.cells[0]
        celda_etiqueta.text = ""
        p = celda_etiqueta.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        corrida = p.add_run(etiqueta + "   ")
        corrida.bold = True
        corrida.font.size = Pt(10)
        corrida.font.color.rgb = AZUL_PRINCIPAL

        celda_valor = fila.cells[1]
        celda_valor.text = ""
        corrida = celda_valor.paragraphs[0].add_run(valor)
        corrida.font.size = Pt(10)

        celda_etiqueta.width = Cm(4.5)
        celda_valor.width = Cm(11)


def indice(documento: Document) -> None:
    """
    Inserta un índice que Word rellena al abrir el archivo.
    El índice se arma con un campo de Word y no con una lista escrita a mano, dado que el campo se actualiza solo con los números de página reales.
    Cabe señalar que esos números dependen de la paginación y, en consecuencia, no se pueden conocer en el momento de generar el archivo.
    Recibe el documento donde escribir.
    """
    documento.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    titulo(documento, "Contenido", nivel=1)

    parrafo(
        documento,
        "Si el índice aparece vacío al abrir el documento, hacer clic derecho "
        "sobre él y elegir la opción de actualizar campo.",
        cursiva=True,
    )

    p = documento.add_paragraph()
    corrida = p.add_run()

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instruccion = OxmlElement("w:instrText")
    instruccion.set(qn("xml:space"), "preserve")
    instruccion.text = 'TOC \\o "1-3" \\h \\z \\u'
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    marcador = OxmlElement("w:t")
    marcador.text = "Hacer clic derecho y actualizar para generar el índice."
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")

    for elemento in (inicio, instruccion, separador, marcador, fin):
        corrida._r.append(elemento)


def pie_de_pagina(documento: Document) -> None:
    """
    Agrega el pie de página con el nombre del proyecto y la numeración de las hojas.
    Recibe el documento sobre el que trabajar y aplica el mismo pie a todas sus secciones.
    El número de página se inserta como campo de Word, de modo que se recalcula solo cuando cambia la paginación.
    """
    for seccion in documento.sections:
        pie = seccion.footer
        p = pie.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        corrida = p.add_run("Pipeline de ventas minoristas    ")
        corrida.font.size = Pt(8)
        corrida.font.color.rgb = GRIS_SUAVE

        corrida = p.add_run()
        inicio = OxmlElement("w:fldChar")
        inicio.set(qn("w:fldCharType"), "begin")
        instruccion = OxmlElement("w:instrText")
        instruccion.text = "PAGE"
        fin = OxmlElement("w:fldChar")
        fin.set(qn("w:fldCharType"), "end")
        for elemento in (inicio, instruccion, fin):
            corrida._r.append(elemento)
        corrida.font.size = Pt(8)
        corrida.font.color.rgb = GRIS_SUAVE


def _leer(ruta: Path, maximo_lineas: int | None = None) -> str:
    """
    Lee un archivo del repositorio para incrustarlo en el documento.
    Recibe la ruta del archivo relativa a la raíz del proyecto y, cuando se indica un máximo de líneas, recorta el contenido a esa cantidad y avisa cuántas quedaron fuera.
    Devuelve el contenido del archivo, o bien un aviso entre corchetes en caso de que no se encuentre.
    """
    completa = RAIZ / ruta
    if not completa.exists():
        return f"[No se encontró el archivo {ruta}]"

    texto = completa.read_text(encoding="utf-8")
    if maximo_lineas is None:
        return texto

    lineas = texto.split("\n")
    if len(lineas) <= maximo_lineas:
        return texto
    return "\n".join(lineas[:maximo_lineas]) + f"\n\n[... {len(lineas) - maximo_lineas} líneas más en el repositorio]"


def _extraer(ruta: Path, desde: str, hasta: str | None = None) -> str:
    """
    Extrae el fragmento de un archivo comprendido entre dos marcas de texto.
    De ese modo se incrusta la parte relevante de un archivo largo sin copiar su contenido, con lo que el documento y el código no se desincronizan.
    Recibe la ruta del archivo, el texto que marca el inicio del fragmento y el texto que marca su final.
    En caso de que se omita la marca final, el fragmento llega hasta el final del archivo.
    Devuelve el fragmento encontrado, o bien un aviso entre corchetes cuando el archivo o la marca de inicio no existen.
    """
    completa = RAIZ / ruta
    if not completa.exists():
        return f"[No se encontró el archivo {ruta}]"

    lineas = completa.read_text(encoding="utf-8").split("\n")
    inicio = next((i for i, linea in enumerate(lineas) if desde in linea), None)
    if inicio is None:
        return f"[No se encontró la marca '{desde}' en {ruta}]"

    if hasta is None:
        return "\n".join(lineas[inicio:])

    fin = next((i for i, linea in enumerate(lineas[inicio + 1 :], inicio + 1) if hasta in linea), len(lineas))
    return "\n".join(lineas[inicio:fin])


def cargar_evidencia() -> dict:
    """
    Lee el reporte de la última corrida para incrustar números reales en el documento.
    Devuelve el reporte completo, o bien un diccionario vacío en caso de que el pipeline todavía no se haya ejecutado.
    """
    ruta = RAIZ / "salida" / "reportes" / "reporte_ultima_corrida.json"
    if not ruta.exists():
        return {}
    return json.loads(ruta.read_text(encoding="utf-8"))


def analizar_argumentos(argumentos: list[str] | None = None) -> argparse.Namespace:
    """
    Define los parámetros de línea de comandos.
    Recibe la lista de argumentos a interpretar y, en caso de que se omita, toma la que dejó el intérprete en sys.argv.
    Devuelve el espacio de nombres con las opciones ya resueltas.
    """
    analizador = argparse.ArgumentParser(
        description="Genera el documento de entrega del proyecto en formato Word"
    )
    analizador.add_argument("--destino", type=Path, default=DESTINO_POR_DEFECTO)
    analizador.add_argument("--autor", default="Jeshua Romero Guadarrama")
    analizador.add_argument(
        "--repositorio",
        default="https://github.com/Jeshua-Romero-Guadarrama/pipeline-ventas-minoristas",
    )
    return analizador.parse_args(argumentos)


def main(argumentos: list[str] | None = None) -> int:
    """
    Construye el documento completo y lo guarda en disco.
    Recibe los argumentos de línea de comandos y devuelve cero cuando el documento se generó correctamente.
    El contenido se importa dentro de la función y no en la cabecera del módulo, puesto que el módulo de contenido importa a su vez las utilidades de formato de este archivo.
    """
    from scripts.contenido_entrega import escribir_contenido

    opciones = analizar_argumentos(argumentos)
    evidencia = cargar_evidencia()

    documento = Document()
    configurar_estilos(documento)

    for seccion in documento.sections:
        seccion.top_margin = Cm(2.2)
        seccion.bottom_margin = Cm(2.2)
        seccion.left_margin = Cm(2.4)
        seccion.right_margin = Cm(2.4)

    metadatos = {
        "autor": opciones.autor,
        "repositorio": opciones.repositorio,
        "fecha": datetime.now().strftime("%d de %B de %Y"),
        "version": "1.0.0",
    }

    portada(documento, metadatos)
    indice(documento)
    escribir_contenido(documento, evidencia, metadatos)
    pie_de_pagina(documento)

    propiedades = documento.core_properties
    propiedades.title = "Pipeline de datos extremo a extremo sobre ventas minoristas"
    propiedades.author = opciones.autor
    propiedades.subject = "Proyecto integrador de ingeniería de datos"
    propiedades.comments = (
        "Documento de entrega final. Incluye arquitectura, infraestructura, "
        "orquestación, transformaciones, pruebas, observabilidad y evidencia."
    )

    opciones.destino.parent.mkdir(parents=True, exist_ok=True)
    documento.save(opciones.destino)

    print(f"Documento generado en {opciones.destino}")
    print(f"Tamaño {opciones.destino.stat().st_size / 1024:.1f} kilobytes")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
